"""Document parsing — PDF, EPUB, DOCX, TXT → structured chapters."""

import hashlib
import os
import re
from html import unescape
from pathlib import Path


def _compute_hash(file_path: str) -> str:
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()[:8]


def _strip_html(html: str) -> str:
    if not html:
        return ""
    text = html
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</(p|div|li|tr|h[1-6])>", "\n\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</(td|th)>", "\t", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    text = unescape(text)
    return text.strip()


def _split_into_chapters(text: str, fallback_title: str = "Document") -> list[dict]:
    chapter_re = re.compile(
        r"^(Chapter|CHAPTER|Part|PART)\s+(\d+|[IVXLC]+)[.:)?\s-]*(.*)?$",
        re.IGNORECASE | re.MULTILINE,
    )
    lines = text.split("\n")
    starts = []

    for i, line in enumerate(lines):
        if chapter_re.match(line.strip()):
            starts.append((i, line.strip()))

    if len(starts) >= 2:
        chapters = []
        for idx, (start_line, title) in enumerate(starts):
            end_line = starts[idx + 1][0] if idx + 1 < len(starts) else len(lines)
            content = "\n".join(lines[start_line:end_line]).strip()
            if content:
                chapters.append({"title": title, "content": content, "index": len(chapters)})

        # Prepend content before first chapter
        if starts[0][0] > 0:
            preface = "\n".join(lines[: starts[0][0]]).strip()
            if len(preface) > 50:
                chapters.insert(0, {"title": "Preface", "content": preface, "index": 0})
                for i, ch in enumerate(chapters):
                    ch["index"] = i

        return chapters

    # Fallback: split into ~3000 char sections
    if len(text) > 6000:
        chapters = []
        paragraphs = text.split("\n\n")
        current = ""
        for para in paragraphs:
            if len(current) + len(para) > 3000 and current:
                chapters.append({
                    "title": f"Section {len(chapters) + 1}",
                    "content": current.strip(),
                    "index": len(chapters),
                })
                current = para
            else:
                current += ("\n\n" if current else "") + para
        if current.strip():
            chapters.append({
                "title": f"Section {len(chapters) + 1}",
                "content": current.strip(),
                "index": len(chapters),
            })
        return chapters

    return [{"title": fallback_title, "content": text, "index": 0}]


def read_file(file_path: str) -> dict:
    path = Path(file_path).resolve()
    ext = path.suffix.lower()
    file_hash = _compute_hash(str(path))

    readers = {
        ".pdf": _read_pdf,
        ".epub": _read_epub,
        ".docx": _read_docx,
    }
    reader = readers.get(ext, _read_text)
    result = reader(str(path))

    result["metadata"]["file_hash"] = file_hash
    result["metadata"]["file_path"] = str(path)
    return result


def _read_text(file_path: str) -> dict:
    text = Path(file_path).read_text(encoding="utf-8", errors="replace")
    name = Path(file_path).stem
    return {
        "metadata": {"title": name, "author": "Unknown", "format": Path(file_path).suffix.lstrip(".")},
        "chapters": _split_into_chapters(text, name),
    }


def _read_pdf(file_path: str) -> dict:
    import fitz  # pymupdf

    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text() + "\n"

    title = doc.metadata.get("title") or Path(file_path).stem
    author = doc.metadata.get("author") or "Unknown"
    doc.close()

    return {
        "metadata": {"title": title, "author": author, "format": "pdf"},
        "chapters": _split_into_chapters(text, title),
    }


def _read_epub(file_path: str) -> dict:
    import ebooklib
    from ebooklib import epub

    book = epub.read_epub(file_path, options={"ignore_ncx": True})

    title = book.get_metadata("DC", "title")
    title = title[0][0] if title else Path(file_path).stem
    creator = book.get_metadata("DC", "creator")
    author = creator[0][0] if creator else "Unknown"

    # Build TOC title map
    toc_titles = {}
    for item in book.toc:
        if isinstance(item, epub.Link):
            # Map filename (without anchor) to title
            href = item.href.split("#")[0]
            toc_titles[href] = item.title

    chapters = []
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        content = item.get_content().decode("utf-8", errors="replace")
        plain = _strip_html(content)
        if plain.strip() and len(plain.strip()) > 20:
            item_href = item.get_name()
            ch_title = toc_titles.get(item_href, f"Chapter {len(chapters) + 1}")
            chapters.append({
                "title": ch_title,
                "content": plain.strip(),
                "index": len(chapters),
            })

    if not chapters:
        chapters.append({"title": title, "content": "Unable to extract content.", "index": 0})

    return {
        "metadata": {"title": title, "author": author, "format": "epub"},
        "chapters": chapters,
    }


def _read_docx(file_path: str) -> dict:
    from docx import Document

    doc = Document(file_path)
    name = Path(file_path).stem
    chapters = []
    current_title = None
    current_content = []

    for para in doc.paragraphs:
        # Heading styles indicate chapter breaks
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            if current_content:
                chapters.append({
                    "title": current_title or f"Section {len(chapters) + 1}",
                    "content": "\n".join(current_content).strip(),
                    "index": len(chapters),
                })
                current_content = []
            current_title = para.text.strip() or f"Chapter {len(chapters) + 1}"
        else:
            if para.text.strip():
                current_content.append(para.text)

    # Last chapter
    if current_content:
        chapters.append({
            "title": current_title or name,
            "content": "\n".join(current_content).strip(),
            "index": len(chapters),
        })

    if not chapters:
        full_text = "\n".join(p.text for p in doc.paragraphs)
        chapters = _split_into_chapters(full_text, name)

    return {
        "metadata": {"title": name, "author": "Unknown", "format": "docx"},
        "chapters": chapters,
    }
